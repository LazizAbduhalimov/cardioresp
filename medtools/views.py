from io import BytesIO

from django.contrib import messages
from django.http import Http404, HttpResponse
from django.shortcuts import get_object_or_404
from django.views.generic import FormView, DeleteView
from docx import Document
from extra_views import UpdateWithInlinesView, CreateWithInlinesView, NamedFormsetsMixin
from django.utils.translation import gettext_lazy as _

from main_app.utils import MenuMixin
from medtools.forms import *
from medtools.models import Patient


class HeartDiseaseToolPage(NamedFormsetsMixin, MenuMixin, CreateWithInlinesView):
    model = Patient
    form_class = PatientCreateForm
    inlines = [EchocardiographyInLine, GeneticResearchInLine, ImmunologicalResearchInLine, BodyMassIndexInLine,
               PilidogramInLine, BiochemicalBloodAnalysisInLine, CoronaryAngiographyInLine, ECGInLine]
    inlines_names = ["Echocardiography", "GeneticResearch", "ImmunologicalResearch", "BodyMassIndex",
                     "Pilidogram", "BiochemicalBloodAnalysis", "CoronaryAngiography", "ECG"]
    template_name = "medtools/heart_disease_tool.html"

    factory_kwargs = {'extra': 1, 'max_num': 1,
                      'can_order': False, 'can_delete': False}

    def get_context_data(self, **kwargs):
        context = super(HeartDiseaseToolPage, self).get_context_data(**kwargs)
        return dict(list(context.items()) + list(self.get_user_context().items()))

    def get_success_url(self):
        return self.object.get_absolute_url()

    def form_valid(self, form):
        form.save()
        self.request.session["patient_id"] = self.object.pk
        return super(HeartDiseaseToolPage, self).form_valid(form)


class HeartDiseaseToolUpdatePage(NamedFormsetsMixin, MenuMixin, UpdateWithInlinesView):
    model = Patient
    form_class = PatientCreateForm
    inlines = [EchocardiographyInLine, GeneticResearchInLine, ImmunologicalResearchInLine, BodyMassIndexInLine,
               PilidogramInLine, BiochemicalBloodAnalysisInLine, CoronaryAngiographyInLine, ECGInLine]
    inlines_names = ["Echocardiography", "GeneticResearch", "ImmunologicalResearch", "BodyMassIndex",
                     "Pilidogram", "BiochemicalBloodAnalysis", "CoronaryAngiography", "ECG"]
    template_name = "medtools/heart_disease_tool_update.html"

    def get_context_data(self, **kwargs):
        context = super(HeartDiseaseToolUpdatePage, self).get_context_data(**kwargs)
        surveys = Survey.objects.all().distinct()
        surveys_ratio_choice = surveys.filter(surveyquestion__has_multiple_choice=False).prefetch_related(
            "surveyquestion_set")
        context["surveys_ratio_choice"] = surveys_ratio_choice
        surveys_multiple_choice = surveys.filter(surveyquestion__has_multiple_choice=True).prefetch_related(
            "surveyquestion_set")
        context["surveys_multiple_choice"] = surveys_multiple_choice
        surveys_result = list()
        patient_id = self.request.session.get("patient_id")
        for survey in surveys_ratio_choice:
            overall_score = survey.get_overall_score(patient_id)
            survey_results_set = survey.surveyresult_set.all()

            if overall_score is None:
                continue

            for result in survey_results_set:
                if result.mark_from <= overall_score <= result.mark_to:
                    surveys_result.append(result.text)
                    break

        for survey in surveys_multiple_choice:
            insd, pri = survey.get_insd_and_pri(patient_id)
            if (pri or insd) is not None:
                context["INSD"] = insd
                context["PRI"] = pri

        context["surveys_result"] = surveys_result

        return dict(list(context.items()) + list(self.get_user_context().items()))

    def get_success_url(self):
        return self.object.get_absolute_url()


class HeartDiseaseToolDeletePage(DeleteView):
    model = Patient
    template_name = "medtools/heart_disease_tool_delete.html"
    success_url = reverse_lazy("heart-disease-tool")


class SurveyView(MenuMixin, FormView):
    template_name = "medtools/survey.html"

    def get_success_url(self):
        question_id = self.request.GET.get("question_id", None)
        url = SurveyQuestion.objects.get(id=question_id).get_next_question()
        if url is None:
            messages.add_message(
                self.request, messages.SUCCESS,
                "Опросник пройден"
            )
            return reverse_lazy("heart-disease-tool-update", kwargs={"pk": self.request.session.get("patient_id")})

        return url

    def get_form_kwargs(self):
        kwargs = super(SurveyView, self).get_form_kwargs()
        question_id = self.request.GET.get("question_id", None)
        kwargs["survey_question"] = SurveyQuestion.objects.get(id=question_id)
        return kwargs

    def get_context_data(self, **kwargs):
        context = super(SurveyView, self).get_context_data(**kwargs)
        question_id = self.request.GET.get("question_id", None)
        if question_id is None:
            raise Http404()

        context["question"] = SurveyQuestion.objects.get(id=question_id)
        context["GET_params"] = f"?question_id={question_id}"
        return dict(list(context.items()) + list(self.get_user_context().items()))


class SurveyMultipleChoicePage(SurveyView):
    form_class = SurveyMultipleAnswerForm

    def form_valid(self, form):
        survey_answer = form.save(commit=False)
        survey_answer.patient = Patient.objects.get(id=self.request.session.get("patient_id"))
        survey_answer.save()
        form.save_m2m()
        return super(SurveyMultipleChoicePage, self).form_valid(form)


class SurveyPage(SurveyView):
    form_class = SurveyAnswerForm

    def form_valid(self, form, *args, **kwargs):
        survey_answer = form.save(commit=False)
        survey_answer.patient = Patient.objects.get(id=self.request.session.get("patient_id"))
        survey_answer.save()
        return super(SurveyPage, self).form_valid(form)


def get_docx_file(request, pk):
    """Return .docx file of the result of the HeartDiseaseToolUpdatePage"""
    patient = get_object_or_404(Patient, pk=pk)
    document = Document()
    docx_title = "result.docx"

    text = str()
    # ----- Analysis data ----- #

    # Main
    text += _("Основная информация") + "\n"
    text += _("Пол") + ": " + patient.get_sex_display() + "\n"
    text += _("Возраст") + ": " + patient.get_age_display() + "\n"
    text += _("Социальный статус") + ": " + patient.get_social_position_display() + "\n"
    text += _("Продолжительность боли") + ": " + patient.get_pain_duration_display() + "\n"
    text += _("Частота сердечных сокращений") + ": " + str(patient.heart_rate) + "\n"
    document.add_paragraph(text)

    # Genetic research
    text = _("Генетическое исследование") + "\n"
    text += "IL-1β 511 T/C (rs16944)" + ": " + patient.geneticresearch.get_IL_1b_511_TC_display() + "\n"
    text += "IL-4 С/Т (rs2243250)" + ": " + patient.geneticresearch.get_IL_4_TC_display() + "\n"
    text += "IL-10 819 C/T (rs1800871)" + ": " + patient.geneticresearch.get_IL_10_819_CT_display() + "\n"
    text += "TNF-α G/A 308 (rs1800629)" + ": " + patient.geneticresearch.get_TNF_a_GA_display() + "\n"
    document.add_paragraph(text)

    # Immunological research
    try:
        text = _("Иммунологическое исследование") + "\n"
        text += "IL-1(b)" + ": " + str(patient.immunologicalresearch.IL_1b) + "\n"
        text += "TNF-a" + ": " + str(patient.immunologicalresearch.TNF_a) + "\n"
        text += "IL-4" + ": " + str(patient.immunologicalresearch.IL_4) + "\n"
        text += "IL-10" + ": " + str(patient.immunologicalresearch.IL_10) + "\n"
        document.add_paragraph(text)
    except ObjectDoesNotExist:
        pass

    # Lipidogram
    try:
        text = _("Липидогамма") + "\n"
        text += _("Общий ХС ммоль/г") + ": " + str(patient.lipidogram.HS) + "\n"
        text += _("ХС ЛПНП. ммоль/г") + ": " + str(patient.lipidogram.HS_LPNP) + "\n"
        text += _("ХС ЛПВП. ммоль/г") + ": " + str(patient.lipidogram.HS_LPVP) + "\n"
        text += _("ТГ. ммоль/г") + ": " + str(patient.lipidogram.TG) + "\n"
        text += _("КА") + ": " + str(patient.lipidogram.KA) + "\n"
        document.add_paragraph(text)
    except ObjectDoesNotExist:
        pass

    # Biochemical blood analysis
    try:
        text = _("Биохимическое исследование крови") + "\n"
        text += _("АЛАТ (U/L)") + ": " + str(patient.biochemicalbloodanalysis.ALAT) + "\n"
        text += _("АСАТ(U/L)") + ": " + str(patient.biochemicalbloodanalysis.ACAT) + "\n"
        text += _("Креатинин (мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.creatinin) + "\n"
        text += _("Мочевина (мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.urea) + "\n"
        text += _("Мочевая кислота(мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.uric_acid) + "\n"
        text += _("Билирубин общий (мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.bilirubin_common) + "\n"
        text += _("Билирубин прямой (мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.bilirubin_direct) + "\n"
        text += _("Билирубин непрямой (мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.bilirubin_indirect) + "\n"
        text += _("Глюкоза (мкмоль/л)") + ": " + str(patient.biochemicalbloodanalysis.glucose) + "\n"
        document.add_paragraph(text)
    except ObjectDoesNotExist:
        pass

    # Body mass index
    try:
        text = _("Индекс массы тела") + "\n"
        text += "Рост (cм)" + ": " + str(patient.bodymassindex.height) + "\n"
        text += "Масса (кг)" + ": " + str(patient.bodymassindex.mass) + "\n"
        document.add_paragraph(text)
    except ObjectDoesNotExist:
        pass

    # Echocardiography
    text = _("Эхокардиография") + "\n"
    text += _("ФВЛЖ (%)") + ": " + str(patient.echocardiography.ejection_fraction) + "\n"
    text += _("КСР (см)") + ": " + str(patient.echocardiography.kcr) + "\n"
    text += _("КДР (см)") + ": " + str(patient.echocardiography.kdr) + "\n"
    text += _("ЗСЛЖ (см)") + ": " + str(patient.echocardiography.eslj) + "\n"
    text += _("МЖП (см)") + ": " + str(patient.echocardiography.mjp) + "\n"
    text += _("ПЖ (см)") + ": " + str(patient.echocardiography.pj) + "\n"
    text += _("МК (см)") + ": " + patient.echocardiography.get_mk_display() + "\n"
    text += _("ЛП (см)") + ": " + str(patient.echocardiography.lp) + "\n"
    document.add_paragraph(text)

    # Coronary Angiography
    text = _("Коронарная ангиография") + "\n"
    text += _("Количество поражений КА") + ": " + patient.coronaryangiography.get_field_display() + "\n\n"
    document.add_paragraph(text)

    text = _("Результаты анализов") + "\n"
    # ----- Result ------ #
    text += "{}, {} \n".format(patient.get_pain_duration_text(), patient.ecg.get_field_display())
    document.add_paragraph(text)

    text = _("Осложнения") + "\n"
    for disease in patient.get_disease_list():
        text += "{} \n".format(disease)

    risk_group = patient.get_risk_group()
    if risk_group is not None:
        text += "{} \n".format(risk_group)

    try:
        mass_index = patient.bodymassindex.get_mass_disease()
        if mass_index is not None:
            text += "{} \n".format(mass_index)
    except ObjectDoesNotExist:
        pass

    try:
        for disease in patient.immunologicalresearch.get_disease_list():
            document.add_paragraph(disease)
    except ObjectDoesNotExist:
        pass

    try:
        for disease in patient.biochemicalbloodanalysis.get_disease_list():
            text += "{} \n".format(disease)
    except ObjectDoesNotExist:
        pass

    try:
        for disease in patient.lipidogram.get_disease_list():
            text += "{} \n".format(disease)
    except ObjectDoesNotExist:
        pass

    document.add_paragraph(text)

    text = _("Показатели ЭХОКГ") + "\n"
    for disease in patient.echocardiography.get_disease_list():
        text += "{} \n".format(disease)
    document.add_paragraph(text)

    text = ""
    surveys = Survey.objects.all().distinct()
    surveys_ratio_choice = surveys.filter(surveyquestion__has_multiple_choice=False).prefetch_related(
        "surveyquestion_set")
    surveys_multiple_choice = surveys.filter(surveyquestion__has_multiple_choice=True).prefetch_related(
        "surveyquestion_set")

    for survey in surveys_ratio_choice:
        overall_score = survey.get_overall_score(pk)
        survey_results_set = survey.surveyresult_set.all()

        if overall_score is None:
            continue

        for result in survey_results_set:
            if result.mark_from <= overall_score <= result.mark_to:
                text += result.text + "\n"
                break

    for survey in surveys_multiple_choice:
        insd, pri = survey.get_insd_and_pri(pk)
        if pri or insd:
            insd_text = _("Индекс числа выбранных дескрипторов")
            pri_text = _("Ранговой индекс боли")
            text += "{}: {}\n".format(insd_text, insd)
            text += "{}: {}\n".format(pri_text, pri)
            continue

    if text != "":
        text = _("Результаты опросников") + "\n" + text

    document.add_paragraph(text)
    document.add_page_break()

    # Prepare document for download
    # -----------------------------
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
